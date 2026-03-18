"""
Development Proposal Generator — AI-powered council proposal creation.
Feature 2: Evidence gathering, Gemini draft, .docx generation, MongoDB storage.
"""
from __future__ import annotations

import io
import json
import math
import logging
import os
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel

from app.core.dependencies import get_current_user
from app.mongodb.models.user import UserMongo
from app.mongodb.models.proposal import ProposalMongo
from app.enums import UserRole
from app.mongodb.database import get_motor_client
from app.core.config import settings

router = APIRouter()
logger = logging.getLogger("proposals")


# ── Auth ──────────────────────────────────────────────────────────────────────

def _require_dev_access(current_user: UserMongo = Depends(get_current_user)) -> UserMongo:
    allowed = {UserRole.COUNCILLOR, UserRole.SUPERVISOR, UserRole.SUPER_ADMIN, "COMMISSIONER"}
    if current_user.role not in allowed:
        raise HTTPException(status_code=403, detail="Councillor/Commissioner access required")
    return current_user


# ── Helpers ───────────────────────────────────────────────────────────────────

def _get_tickets_col():
    client = get_motor_client()
    uri = settings.MONGODB_URI
    db_name = uri.rsplit("/", 1)[-1].split("?")[0] or "civicai"
    if not db_name or db_name.startswith("mongodb"):
        db_name = "civicai"
    return client[db_name]["tickets"]


def haversine_km(lat1: float, lng1: float, lat2: float, lng2: float) -> float:
    """Haversine distance in kilometers."""
    R = 6371
    dlat = math.radians(lat2 - lat1)
    dlng = math.radians(lng2 - lng1)
    a = (math.sin(dlat / 2) ** 2
         + math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.sin(dlng / 2) ** 2)
    return R * 2 * math.asin(math.sqrt(a))


# ── Budget lookup ─────────────────────────────────────────────────────────────

BUDGET_TABLE: Dict[str, Dict] = {
    "road_resurfacing": {
        "label": "Road Resurfacing",
        "unit": "per km",
        "cost_inr": 2_500_000,
        "default_scope": "0.5 km stretch",
        "default_total": 1_250_000,
    },
    "streetlight_installation": {
        "label": "Streetlight Installation",
        "unit": "per light",
        "cost_inr": 25_000,
        "default_scope": "20 streetlights",
        "default_total": 500_000,
    },
    "drainage_improvement": {
        "label": "Drainage Improvement",
        "unit": "per 100m drain",
        "cost_inr": 800_000,
        "default_scope": "200m drain",
        "default_total": 1_600_000,
    },
    "park_open_space": {
        "label": "Park / Open Space Development",
        "unit": "per 1000 sqm",
        "cost_inr": 1_500_000,
        "default_scope": "2000 sqm park",
        "default_total": 3_000_000,
    },
    "water_pipeline": {
        "label": "Water Pipeline",
        "unit": "per km",
        "cost_inr": 3_000_000,
        "default_scope": "0.3 km pipeline",
        "default_total": 900_000,
    },
    "waste_collection_point": {
        "label": "Waste Collection Point",
        "unit": "per unit",
        "cost_inr": 150_000,
        "default_scope": "3 collection points",
        "default_total": 450_000,
    },
    "community_center": {
        "label": "Community Centre",
        "unit": "per sqm construction",
        "cost_inr": 18_000,
        "default_scope": "500 sqm building",
        "default_total": 9_000_000,
    },
}


# ── Request schema ─────────────────────────────────────────────────────────────

class GenerateProposalRequest(BaseModel):
    ward_id: Optional[int] = None
    zone_cell_id: Optional[str] = None
    zone_lat: float
    zone_lng: float
    development_type: str
    estimated_cost: Optional[float] = None
    councillor_name: str
    ward_name: str
    additional_context: Optional[str] = ""


# ── Gemini helper ──────────────────────────────────────────────────────────────

async def _call_gemini_proposal(prompt: str) -> Dict:
    """Call Gemini for proposal JSON. Retries once on parse failure."""
    from langchain_google_genai import ChatGoogleGenerativeAI

    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        google_api_key=settings.GEMINI_API_KEY,
        temperature=0.2,
    )

    for attempt in range(2):
        try:
            resp = await llm.ainvoke([("user", prompt)])
            text = resp.content.strip()
            if text.startswith("```"):
                lines = text.split("\n")
                text = "\n".join(lines[1:-1]) if len(lines) > 2 else text
            return json.loads(text)
        except Exception as e:
            logger.warning(f"Gemini proposal attempt {attempt + 1} failed: {e}")
            if attempt == 1:
                raise HTTPException(
                    status_code=500,
                    detail="AI proposal generation failed. Please try again in a moment.",
                )
    return {}


# ── docx builder ──────────────────────────────────────────────────────────────

def _build_docx(proposal: Dict, ward_name: str, councillor_name: str, today: str) -> str:
    """Build a .docx from the AI proposal JSON. Returns the file path."""
    try:
        from docx import Document
        from docx.shared import Pt, Rgb, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed on the server. Run: pip install python-docx",
        )

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    DARK_BLUE = RGBColor(0, 51, 102)

    def add_section_heading(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        # Bottom border on the paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "003366")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_body(text: str):
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(11)
        return p

    # ── Title block ────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("WARD DEVELOPMENT PROPOSAL")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = DARK_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(proposal.get("title", ""))
    sub_run.font.size = Pt(13)

    doc.add_paragraph()  # spacer

    # ── Metadata table ─────────────────────────────────────────────────────────
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    rows_data = [
        ("Reference:", proposal.get("reference_number_placeholder", "")),
        ("Ward:", ward_name),
        ("Councillor:", councillor_name),
        ("Date:", today),
    ]
    for i, (label, value) in enumerate(rows_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            # Remove borders for clean look
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    doc.add_paragraph()

    # ── Sections ───────────────────────────────────────────────────────────────
    text_sections = [
        ("Executive Summary", "executive_summary"),
        ("Problem Statement", "problem_statement"),
        ("Proposed Solution", "proposed_solution"),
        ("Location Justification", "location_justification"),
        ("Beneficiary Analysis", "beneficiary_analysis"),
    ]
    for heading, key in text_sections:
        add_section_heading(heading)
        add_body(proposal.get(key, ""))

    # ── Budget table ───────────────────────────────────────────────────────────
    budget_items = proposal.get("budget_breakdown", [])
    if budget_items:
        add_section_heading("Budget Breakdown")
        budget_table = doc.add_table(rows=1 + len(budget_items), cols=4)
        budget_table.style = "Table Grid"

        # Header row
        headers = ["Item", "Quantity", "Unit Cost", "Total"]
        for j, h in enumerate(headers):
            cell = budget_table.rows[0].cells[j]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            # Blue background
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "003366")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)

        for i, item in enumerate(budget_items):
            row = budget_table.rows[i + 1]
            row.cells[0].text = item.get("item", "")
            row.cells[1].text = str(item.get("quantity", ""))
            row.cells[2].text = str(item.get("unit_cost", ""))
            row.cells[3].text = str(item.get("total", ""))
            # Alternating row shading
            if i % 2 == 1:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "EBF0FF")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

        tp = doc.add_paragraph()
        run = tp.add_run(f"Total Budget Requested: {proposal.get('total_budget_requested', '')}")
        run.bold = True
        run.font.size = Pt(11)

    # ── Expected outcomes ──────────────────────────────────────────────────────
    outcomes = proposal.get("expected_outcomes", [])
    if outcomes:
        add_section_heading("Expected Outcomes")
        for outcome in outcomes:
            doc.add_paragraph(outcome, style="List Bullet")

    # ── Timeline table ─────────────────────────────────────────────────────────
    timeline = proposal.get("implementation_timeline", [])
    if timeline:
        add_section_heading("Implementation Timeline")
        tl_table = doc.add_table(rows=1 + len(timeline), cols=3)
        tl_table.style = "Table Grid"
        for j, h in enumerate(["Phase", "Activity", "Duration"]):
            cell = tl_table.rows[0].cells[j]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "003366")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
        for i, phase in enumerate(timeline):
            row = tl_table.rows[i + 1]
            row.cells[0].text = phase.get("phase", "")
            row.cells[1].text = phase.get("activity", "")
            row.cells[2].text = phase.get("duration", "")
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

    # ── Recommendation ─────────────────────────────────────────────────────────
    add_section_heading("Recommendation")
    rec_p = doc.add_paragraph(proposal.get("recommendation", ""))
    for run in rec_p.runs:
        run.italic = True
        run.font.size = Pt(11)

    # ── Footer ─────────────────────────────────────────────────────────────────
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"{ward_name} Ward  |  JanVedha AI Generated  |  {today}  |  Costs are estimates subject to survey"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Save
    os.makedirs("/tmp/proposals", exist_ok=True)
    proposal_id = uuid4().hex[:8]
    path = f"/tmp/proposals/{proposal_id}.docx"
    doc.save(path)
    return path, proposal_id


# ── Endpoints ──────────────────────────────────────────────────────────────────

@router.post("/generate")
async def generate_proposal(
    req: GenerateProposalRequest,
    current_user: UserMongo = Depends(_require_dev_access),
):
    effective_ward = req.ward_id or current_user.ward_id

    # Validate development type
    if req.development_type not in BUDGET_TABLE:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid development_type. Choose from: {', '.join(BUDGET_TABLE.keys())}",
        )

    budget_info = BUDGET_TABLE[req.development_type]

    # ── Step 1: Gather zone evidence ─────────────────────────────────────────
    col = _get_tickets_col()
    since_12m = datetime.utcnow() - timedelta(days=365)

    raw_tickets = await col.find(
        {
            "ward_id": effective_ward,
        }
    ).to_list(length=5000)

    raw_tickets = [t for t in raw_tickets if isinstance(t.get("created_at"), datetime) and t["created_at"].replace(tzinfo=None) >= since_12m]


    # Filter by Haversine distance ≤ 1km
    zone_tickets = []
    for t in raw_tickets:
        loc = t.get("location")
        if loc and isinstance(loc, dict):
            coords = loc.get("coordinates")
            if coords and len(coords) >= 2:
                t_lng, t_lat = coords[0], coords[1]
                dist = haversine_km(req.zone_lat, req.zone_lng, t_lat, t_lng)
                if dist <= 1.0:
                    zone_tickets.append(t)

    total_complaints = len(zone_tickets)

    # Warn if very low data
    low_data_warning = None
    if total_complaints < 3:
        low_data_warning = (
            f"Low data confidence — only {total_complaints} complaint(s) found "
            "in this area. The proposal will have limited evidence base."
        )

    # Evidence aggregation
    from collections import Counter
    cat_counter = Counter()
    unresolved = 0
    resolution_times = []
    oldest_unresolved: Optional[datetime] = None
    most_recent_date: Optional[datetime] = None

    now = datetime.utcnow()

    for t in zone_tickets:
        cat = (t.get("issue_category") or "other").lower()
        cat_counter[cat] += 1

        created_at = t.get("created_at")
        if isinstance(created_at, datetime):
            if most_recent_date is None or created_at > most_recent_date:
                most_recent_date = created_at

        status = (t.get("status") or "OPEN").upper()
        if status not in ("CLOSED", "RESOLVED", "REJECTED", "CLOSED_UNVERIFIED"):
            unresolved += 1
            if isinstance(created_at, datetime):
                if oldest_unresolved is None or created_at < oldest_unresolved:
                    oldest_unresolved = created_at
        else:
            resolved_at = t.get("resolved_at")
            if isinstance(created_at, datetime) and isinstance(resolved_at, datetime):
                days_to_resolve = (resolved_at - created_at).days
                resolution_times.append(days_to_resolve)

    avg_resolution_days = (
        round(sum(resolution_times) / len(resolution_times), 1) if resolution_times else "N/A"
    )
    oldest_unresolved_str = (
        oldest_unresolved.strftime("%d %b %Y") if oldest_unresolved else "N/A"
    )

    # ── Step 2: Budget ────────────────────────────────────────────────────────
    if req.estimated_cost and req.estimated_cost > 0:
        estimated_cost = req.estimated_cost
        cost_note = "as provided by councillor"
    else:
        estimated_cost = budget_info["default_total"]
        cost_note = f"approximate estimate based on {budget_info['default_scope']}"

    today_date = datetime.utcnow().strftime("%d %B %Y")

    # ── Step 3: Gemini proposal ───────────────────────────────────────────────
    prompt = f"""You are drafting a formal infrastructure development proposal for a Ward Councillor to present at a Municipal Corporation council meeting in India.

Generate a complete proposal with the following sections. Use ONLY the data provided below — do not invent statistics. Write in formal English suitable for an official government document.

INPUT DATA:
- Ward: {req.ward_name}
- Councillor: {req.councillor_name}
- Development type: {budget_info['label']}
- Zone location: approximately {req.zone_lat}°N, {req.zone_lng}°E
- Analysis radius: 1km from zone center
- Ticket evidence (past 12 months):
    Total complaints in area: {total_complaints}
    Complaints by category: {dict(cat_counter)}
    Currently unresolved: {unresolved}
    Average resolution time: {avg_resolution_days} days
    Oldest unresolved complaint filed: {oldest_unresolved_str}
- Estimated cost: ₹{estimated_cost:,.0f} ({cost_note})
- Estimated scope: {budget_info['default_scope']}
- Additional context from councillor: {req.additional_context or 'None provided'}

Generate the proposal in this exact JSON structure:
{{
  "title": "Proposal title (formal, specific)",
  "reference_number_placeholder": "WC/{req.ward_name}/2025/XXX",
  "date": "{today_date}",
  "executive_summary": "2-3 sentences summarizing the need and proposed solution",
  "problem_statement": "2-3 paragraphs describing the infrastructure gap, using the ticket data as evidence. Mention specific numbers: complaint count, resolution failure, duration of the problem.",
  "proposed_solution": "1-2 paragraphs describing what will be built/done, where, and at what scale.",
  "location_justification": "1 paragraph explaining why this specific location was selected based on the data.",
  "beneficiary_analysis": "Estimated number of residents who will benefit. Use complaint density to infer approximate population served. Note this is an estimate.",
  "budget_breakdown": [
    {{"item": "...", "quantity": "...", "unit_cost": "₹...", "total": "₹..."}}
  ],
  "total_budget_requested": "₹{estimated_cost:,.0f}",
  "expected_outcomes": [
    "Outcome 1 (measurable where possible)",
    "Outcome 2",
    "Outcome 3"
  ],
  "implementation_timeline": [
    {{"phase": "Phase 1", "activity": "...", "duration": "..."}},
    {{"phase": "Phase 2", "activity": "...", "duration": "..."}},
    {{"phase": "Phase 3", "activity": "...", "duration": "..."}}
  ],
  "recommendation": "1 paragraph formal recommendation statement requesting council approval."
}}

Cost estimate note to include: Cost estimate is approximate. Subject to detailed survey and tender process.
Respond with ONLY the JSON object. No preamble, no markdown fences."""

    ai_proposal = await _call_gemini_proposal(prompt)
    ai_proposal["total_budget_requested"] = f"₹{estimated_cost:,.0f}"

    # ── Step 4: Build .docx ────────────────────────────────────────────────────
    try:
        docx_path, proposal_id = _build_docx(ai_proposal, req.ward_name, req.councillor_name, today_date)
    except Exception as e:
        logger.error(f"docx build failed: {e}")
        # Continue without docx — still return the JSON
        docx_path = None
        proposal_id = uuid4().hex[:8]

    # ── Step 5: Persist to MongoDB ─────────────────────────────────────────────
    proposal_doc = ProposalMongo(
        proposal_id=proposal_id,
        ward_id=effective_ward,
        ward_name=req.ward_name,
        councillor_name=req.councillor_name,
        development_type=req.development_type,
        zone_lat=req.zone_lat,
        zone_lng=req.zone_lng,
        zone_cell_id=req.zone_cell_id,
        estimated_cost=estimated_cost,
        total_complaints_evidence=total_complaints,
        ai_proposal_json=ai_proposal,
        docx_path=docx_path,
        status="draft",
    )
    await proposal_doc.insert()

    return {
        "proposal_id": proposal_id,
        "title": ai_proposal.get("title", ""),
        "proposal_json": ai_proposal,
        "download_url": f"/api/proposals/{proposal_id}/download",
        "created_at": proposal_doc.created_at.isoformat(),
        "total_complaints_evidence": total_complaints,
        "estimated_cost": estimated_cost,
        "low_data_warning": low_data_warning,
        "evidence": {
            "total_complaints": total_complaints,
            "complaints_by_category": dict(cat_counter),
            "unresolved_count": unresolved,
            "avg_resolution_days": avg_resolution_days,
            "oldest_unresolved_date": oldest_unresolved_str,
            "most_recent_complaint_date": most_recent_date.strftime("%d %b %Y") if most_recent_date else "N/A",
        },
    }


@router.get("/{proposal_id}/download")
async def download_proposal(
    proposal_id: str,
    current_user: UserMongo = Depends(_require_dev_access),
):
    """Download the .docx for a proposal."""
    doc = await ProposalMongo.find_one(ProposalMongo.proposal_id == proposal_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Proposal not found")

    if not doc.docx_path or not os.path.exists(doc.docx_path):
        raise HTTPException(
            status_code=404,
            detail="Document file not found. It may have been generated without python-docx support.",
        )

    filename = f"Ward_Proposal_{doc.development_type}_{proposal_id}.docx"
    return FileResponse(
        path=doc.docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


@router.get("")
async def list_proposals(
    ward_id: Optional[int] = Query(None),
    current_user: UserMongo = Depends(_require_dev_access),
):
    """List past proposals for a ward (latest 20)."""
    effective_ward = ward_id or current_user.ward_id

    docs = await ProposalMongo.find(
        ProposalMongo.ward_id == effective_ward
    ).sort(-ProposalMongo.created_at).limit(20).to_list()

    return [
        {
            "proposal_id": d.proposal_id,
            "title": d.ai_proposal_json.get("title", "Untitled"),
            "development_type": d.development_type,
            "estimated_cost": d.estimated_cost,
            "total_complaints_evidence": d.total_complaints_evidence,
            "created_at": d.created_at.isoformat(),
            "status": d.status,
        }
        for d in docs
    ]


@router.get("/{proposal_id}")
async def get_proposal(
    proposal_id: str,
    current_user: UserMongo = Depends(_require_dev_access),
):
    """Get a specific proposal by ID."""
    doc = await ProposalMongo.find_one(ProposalMongo.proposal_id == proposal_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Proposal not found")

    return {
        "proposal_id": doc.proposal_id,
        "title": doc.ai_proposal_json.get("title", "Untitled"),
        "development_type": doc.development_type,
        "estimated_cost": doc.estimated_cost,
        "total_complaints_evidence": doc.total_complaints_evidence,
        "proposal_json": doc.ai_proposal_json,
        "download_url": f"/api/proposals/{doc.proposal_id}/download",
        "created_at": doc.created_at.isoformat(),
        "status": doc.status,
        "zone_lat": doc.zone_lat,
        "zone_lng": doc.zone_lng,
        "ward_name": doc.ward_name,
        "councillor_name": doc.councillor_name,
    }
